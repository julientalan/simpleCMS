"""
app.py — Standalone content editor tool (Flask).

Two-step flow:
Step 1: Dark, polished landing screen ("Content Capture") with a discreet small
        logo/brand line, title, subtitle, and a glassmorphism card containing
        the URL field + Load button, over a subtle grid + glow background.
Step 2: Full-page preview only. Two buttons: "Work with AI" (downloads a prompt,
        then turns into "Import AI Result") and "Finished - Save Changes"
        (downloads Word + HTML, shows a clear confirmation modal with retry buttons
        and a strong warning not to close the page before both files are saved).
Short explanatory popups appear automatically when entering step 2, and briefly
when clicking "Work with AI", so the user always knows what to do next.
"""

import io
from urllib.parse import urljoin

from flask import Flask, request, send_file
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.section import WD_ORIENT

app = Flask(__name__)

HARD_REMOVE_TAGS = {"script"}
SOFT_IGNORE_TAGS = {"footer", "nav", "form", "button"}
IGNORED_CLASS_ID_HINTS = ("nav", "menu", "footer", "cookie", "sidebar", "breadcrumb", "site-header", "topbar")

STEP1_SHELL = """<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>Content Capture</title>
<link rel="icon" type="image/png" href="https://www.julienrio.com/images/logo.png">
<style>
*{{box-sizing:border-box;}}
html,body{{height:100%;margin:0;}}
body{{
  font-family:'Inter','Segoe UI',Arial,sans-serif;
  background:#0b0c17;
  display:flex;align-items:center;justify-content:center;
  color:#fff;
  overflow:hidden;
  position:relative;
}}
.__glow_a__{{position:absolute;top:-180px;left:-120px;width:420px;height:420px;border-radius:50%;
  background:radial-gradient(circle, rgba(230,0,126,.35), transparent 70%);filter:blur(10px);}}
.__glow_b__{{position:absolute;bottom:-200px;right:-140px;width:480px;height:480px;border-radius:50%;
  background:radial-gradient(circle, rgba(0,114,206,.30), transparent 70%);filter:blur(10px);}}
.__grid__{{position:absolute;inset:0;background-image:
  linear-gradient(rgba(255,255,255,.03) 1px, transparent 1px),
  linear-gradient(90deg, rgba(255,255,255,.03) 1px, transparent 1px);
  background-size:48px 48px;}}

.__wrap__{{position:relative;z-index:2;text-align:center;width:100%;max-width:520px;padding:0 24px;}}

.__brand__{{display:flex;align-items:center;justify-content:center;gap:8px;margin-bottom:36px;opacity:.75;}}
.__brand__ img{{height:20px;width:20px;border-radius:5px;}}
.__brand__ span{{font-size:12px;letter-spacing:.06em;text-transform:uppercase;color:#9a9db8;}}

.__title__{{font-size:30px;font-weight:700;margin:0 0 10px;color:#fff;letter-spacing:-.01em;}}
.__subtitle__{{font-size:15px;color:#9a9db8;margin:0 0 40px;line-height:1.5;}}

.__card__{{
  background:rgba(255,255,255,.045);
  border:1px solid rgba(255,255,255,.09);
  border-radius:16px;
  padding:8px;
  display:flex;gap:8px;
  box-shadow:0 20px 60px rgba(0,0,0,.45);
  backdrop-filter:blur(10px);
}}
input[type=text]{{
  flex:1;
  padding:15px 18px;
  border-radius:11px;
  border:none;
  font-size:14px;
  background:rgba(255,255,255,.05);
  color:#fff;
  outline:none;
  font-family:inherit;
}}
input[type=text]::placeholder{{color:#6f7290;}}
input[type=text]:focus{{background:rgba(255,255,255,.09);}}
button{{
  border:none;
  padding:15px 28px;
  border-radius:11px;
  cursor:pointer;
  font-size:14px;
  font-weight:600;
  color:#fff;
  background:linear-gradient(135deg,#e6007e,#8f0060);
  transition:opacity .15s, transform .15s;
  font-family:inherit;
  white-space:nowrap;
}}
button:hover{{opacity:.92;transform:translateY(-1px);}}

.__footnote__{{margin-top:22px;font-size:12.5px;color:#5c5f7a;}}
</style>
</head>
<body>
<div class="__grid__"></div>
<div class="__glow_a__"></div>
<div class="__glow_b__"></div>

<div class="__wrap__">
  <div class="__brand__">
    <img src="https://www.julienrio.com/images/logo.png" alt="logo">
    <span>Julien Rio</span>
  </div>
  <p class="__title__">Content Capture</p>
  <p class="__subtitle__">Paste a web page URL below to load its content, edit it directly<br>or rework it with your favorite AI.</p>
  <div class="__card__">
    <input type="text" id="url_input" placeholder="https://your-site.com/your-page">
    <button onclick="loadPage()">Load page</button>
  </div>
  <p class="__footnote__">Works with any public web page &middot; No account needed</p>
</div>

<script>
function loadPage(){{
  var url = document.getElementById('url_input').value;
  if(!url) return;
  window.location = '/?url=' + encodeURIComponent(url) + '&dynamic=false';
}}
document.getElementById('url_input').addEventListener('keydown', function(e){{
  if(e.key === 'Enter') loadPage();
}});
</script>
</body>
</html>"""

STEP2_SHELL = """<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>Content Editor</title>
<link rel="icon" type="image/png" href="https://www.julienrio.com/images/logo.png">
<style>
html,body{{height:100%;margin:0;}}
body{{font-family:'Segoe UI',Arial,sans-serif;background:#f4f5f7;}}
.__app_bar__{{background:#1a1a2e;padding:12px 20px;display:flex;gap:10px;align-items:center;justify-content:flex-end;
box-shadow:0 2px 8px rgba(0,0,0,.15);}}
.__app_bar__ button{{border:none;padding:10px 18px;border-radius:8px;cursor:pointer;font-size:14px;font-weight:600;
color:#fff;transition:opacity .15s;}}
.__app_bar__ button:hover{{opacity:.85;}}
.btn-ai{{background:#0072ce;}}
.btn-finish{{background:#16a34a;}}
.__frame_wrap__{{padding:16px;}}
iframe{{width:100%;height:85vh;border:1px solid #ddd;border-radius:8px;background:#fff;}}

.__overlay__{{display:none;position:fixed;inset:0;background:rgba(0,0,0,.45);align-items:center;justify-content:center;z-index:50;}}
.__modal__{{background:#fff;border-radius:10px;max-width:480px;width:90%;padding:24px;box-shadow:0 8px 30px rgba(0,0,0,.25);}}
.__modal__ h3{{margin:0 0 12px;font-size:17px;}}
.__modal__ p{{font-size:14px;line-height:1.5;color:#333;margin:0 0 12px;}}
.__modal__ .warn{{background:#fff3f3;border:1px solid #f5a3a3;color:#7a1f1f;padding:10px 12px;border-radius:8px;
font-size:13px;font-weight:600;margin-bottom:14px;}}
.__modal__ .row{{display:flex;gap:8px;flex-wrap:wrap;margin-top:8px;}}
.__modal__ button{{border:none;padding:9px 14px;border-radius:7px;cursor:pointer;font-size:13px;font-weight:600;color:#fff;}}
.btn-primary{{background:#0072ce;}}
.btn-secondary{{background:#666;}}
.btn-retry{{background:#e6a700;}}

.__error_box__{{display:none;margin:0 20px 12px;padding:14px;background:#fff3f3;border:1px solid #f5a3a3;
border-radius:8px;color:#7a1f1f;font-size:13px;white-space:pre-wrap;font-family:monospace;}}
.__error_box__ button{{margin-top:8px;background:#7a1f1f;color:#fff;border:none;padding:6px 12px;
border-radius:6px;cursor:pointer;font-size:12px;}}
</style>
</head>
<body>
<div class="__app_bar__">
  <button class="btn-ai" id="ai_button" onclick="workWithAi()">Work with AI</button>
  <input type="file" id="import_input" accept="application/json" style="display:none" onchange="importJson(event)">
  <button class="btn-finish" onclick="finishEditing()">Finished - Save Changes</button>
</div>
<div class="__error_box__" id="error_box">
  <div id="error_text"></div>
  <button onclick="copyError()">Copy error</button>
</div>
<div class="__frame_wrap__">
  <iframe id="content_frame" srcdoc="{iframe_srcdoc}"></iframe>
</div>

<div class="__overlay__" id="intro_overlay">
  <div class="__modal__">
    <h3>How this works</h3>
    <p>Click any text on the page (including titles) to edit it directly. Or click <strong>Work with AI</strong> to rewrite the content with your favorite AI assistant. When you're done, click <strong>Finished - Save Changes</strong> to download your update.</p>
    <div class="row"><button class="btn-primary" onclick="closeOverlay('intro_overlay')">Got it</button></div>
  </div>
</div>

<div class="__overlay__" id="ai_overlay">
  <div class="__modal__">
    <h3>Working with AI</h3>
    <p>A file named <strong>ai_prompt.txt</strong> was just downloaded. Open it, add your instructions where indicated, and paste the whole content into your favorite AI (Perplexity, Claude, Mistral, ChatGPT...).</p>
    <p>Ask the AI to give you back the final result as a downloadable <strong>.json</strong> file. Once you have it, click <strong>Import AI Result</strong> below to bring it back into the page.</p>
    <div class="row"><button class="btn-primary" onclick="closeOverlay('ai_overlay')">Got it</button></div>
  </div>
</div>

<div class="__overlay__" id="finish_overlay">
  <div class="__modal__">
    <h3>Your update is ready</h3>
    <p>Two files have just been downloaded: <strong>content_update.docx</strong> and <strong>edited_page.html</strong>.</p>
    <p>Please send <strong>both files</strong> to the web team so they can apply the update.</p>
    <div class="warn">Important: make sure both files have downloaded successfully BEFORE closing this page. If you close it now, your changes will be lost.</div>
    <div class="row">
      <button class="btn-retry" onclick="downloadWord()">Retry Word download</button>
      <button class="btn-retry" onclick="downloadHtml()">Retry HTML download</button>
    </div>
    <div class="row"><button class="btn-secondary" onclick="closeOverlay('finish_overlay')">Close</button></div>
  </div>
</div>

<script>
var SOURCE_URL = "{url_value}";
var aiPromptDownloaded = false;

function closeOverlay(id){{ document.getElementById(id).style.display = 'none'; }}
function openOverlay(id){{ document.getElementById(id).style.display = 'flex'; }}

window.addEventListener('load', function(){{ openOverlay('intro_overlay'); }});

function getFrameParts(){{
  var iframeEl = document.getElementById('content_frame');
  var win = iframeEl.contentWindow;
  var doc = win.document;
  var originals = win.__ORIGINALS__ || {{}};
  return {{win: win, doc: doc, originals: originals}};
}}

function collectBlocks(){{
  var parts = getFrameParts();
  var blocks = [];
  parts.doc.querySelectorAll('[data-edit-id]').forEach(function(el){{
    var id = el.getAttribute('data-edit-id');
    blocks.push({{id: id, text: el.textContent}});
  }});
  return blocks;
}}

function showError(msg){{
  document.getElementById('error_text').textContent = msg;
  document.getElementById('error_box').style.display = 'block';
}}
function hideError(){{ document.getElementById('error_box').style.display = 'none'; }}
function copyError(){{
  navigator.clipboard.writeText(document.getElementById('error_text').textContent);
}}

function downloadPromptFile(){{
  var blocks = collectBlocks();
  var jsonPayload = JSON.stringify({{ source_url: SOURCE_URL, blocks: blocks }}, null, 2);
  var prompt = [
    "You are helping me rewrite the text content of a web page.",
    "",
    "INSTRUCTIONS:",
    "1. The reference page (what it currently looks like) is here: " + SOURCE_URL,
    "2. Below is a JSON object listing every editable text block on that page, each with a unique 'id' and its current 'text'.",
    "3. Rewrite/improve the 'text' values only, based on my instructions (see below).",
    "4. Do NOT change, remove, add, or reorder any 'id'. Keep the exact same number of blocks, in the exact same order.",
    "5. Do NOT add new keys or blocks. Only edit the 'text' value inside each block.",
    "6. Provide the final result as a downloadable .json file (not pasted as plain text in the chat), keeping the exact same structure (source_url + blocks array with id/text). If your interface cannot generate downloadable files, then reply with ONLY the final JSON object, no explanation, no markdown code fence, so I can copy it into a text file myself and save it with a .json extension.",
    "",
    "MY INSTRUCTIONS FOR THE REWRITE: [describe here what you want changed - tone, length, SEO, etc.]",
    "",
    "JSON TO EDIT:",
    jsonPayload
  ].join("\\n");
  var blob = new Blob([prompt], {{type: 'text/plain'}});
  var a = document.createElement('a');
  a.href = URL.createObjectURL(blob);
  a.download = 'ai_prompt.txt';
  a.click();
}}

function workWithAi(){{
  hideError();
  downloadPromptFile();
  aiPromptDownloaded = true;
  var btn = document.getElementById('ai_button');
  btn.textContent = 'Import AI Result';
  btn.onclick = function(){{ document.getElementById('import_input').click(); }};
  openOverlay('ai_overlay');
}}

function importJson(event){{
  hideError();
  var file = event.target.files[0];
  if(!file) return;
  var reader = new FileReader();
  reader.onload = function(e){{
    var parsed;
    try {{ parsed = JSON.parse(e.target.result); }}
    catch(err){{
      showError("The file is not valid JSON.\\n\\nDetails: " + err.message + "\\n\\nAsk your AI to fix the JSON syntax and return ONLY a valid JSON object (no markdown code fences, no extra text) with the same structure (source_url + blocks array with id/text).");
      return;
    }}
    if(!parsed || !Array.isArray(parsed.blocks)){{
      showError("The JSON is missing a 'blocks' array.\\n\\nAsk your AI to return the JSON in this exact structure:\\n{{\\n  \\"source_url\\": \\"...\\",\\n  \\"blocks\\": [ {{ \\"id\\": \\"t0\\", \\"text\\": \\"...\\" }}, ... ]\\n}}");
      return;
    }}
    var parts = getFrameParts();
    var currentIds = [];
    parts.doc.querySelectorAll('[data-edit-id]').forEach(function(el){{ currentIds.push(el.getAttribute('data-edit-id')); }});
    var newIds = parsed.blocks.map(function(b){{ return b.id; }});
    var missing = currentIds.filter(function(id){{ return newIds.indexOf(id) === -1; }});
    var extra = newIds.filter(function(id){{ return currentIds.indexOf(id) === -1; }});
    var duplicates = newIds.filter(function(id, idx){{ return newIds.indexOf(id) !== idx; }});
    if(missing.length || extra.length || duplicates.length || newIds.length !== currentIds.length){{
      var msg = "The imported JSON does not match the original structure.\\n\\n";
      msg += "Expected " + currentIds.length + " blocks, got " + newIds.length + ".\\n";
      if(missing.length) msg += "\\nMissing ids (must be present, unchanged): " + missing.join(", ");
      if(extra.length) msg += "\\nUnexpected ids (must be removed, not part of the original): " + extra.join(", ");
      if(duplicates.length) msg += "\\nDuplicate ids (must be unique): " + duplicates.join(", ");
      msg += "\\n\\nAsk your AI: 'Please fix the JSON so it contains exactly these ids, once each, in this order: " + currentIds.join(", ") + ". Do not add, remove, or duplicate any id. Only edit the text values.'";
      showError(msg);
      return;
    }}
    var textById = {{}};
    parsed.blocks.forEach(function(b){{ textById[b.id] = b.text; }});
    parts.doc.querySelectorAll('[data-edit-id]').forEach(function(el){{
      var id = el.getAttribute('data-edit-id');
      if(id in textById){{ el.textContent = textById[id]; }}
    }});
    hideError();
    var btn = document.getElementById('ai_button');
    btn.textContent = 'Work with AI';
    btn.onclick = workWithAi;
  }};
  reader.readAsText(file);
  event.target.value = "";
}}

function downloadWord(){{
  var parts = getFrameParts();
  var data = [];
  parts.doc.querySelectorAll('[data-edit-id]').forEach(function(el){{
    var id = el.getAttribute('data-edit-id');
    var current = el.textContent;
    var original = (id in parts.originals) ? parts.originals[id] : current;
    data.push({{original: original, current: current}});
  }});
  return fetch('/export/docx', {{
    method: 'POST',
    headers: {{'Content-Type': 'application/json'}},
    body: JSON.stringify({{blocks: data}})
  }}).then(function(resp){{ return resp.blob(); }})
    .then(function(blob){{
      var a = document.createElement('a');
      a.href = URL.createObjectURL(blob);
      a.download = 'content_update.docx';
      a.click();
    }});
}}

function downloadHtml(){{
  var parts = getFrameParts();
  var clone = parts.doc.documentElement.cloneNode(true);
  clone.querySelectorAll('[data-edit-id]').forEach(function(el){{
    el.removeAttribute('data-edit-id');
    el.removeAttribute('contenteditable');
    el.removeAttribute('style');
  }});
  clone.querySelectorAll('script').forEach(function(s){{ s.remove(); }});
  var htmlStr = '<!doctype html>\\n' + clone.outerHTML;
  var blob = new Blob([htmlStr], {{type: 'text/html'}});
  var a = document.createElement('a');
  a.href = URL.createObjectURL(blob);
  a.download = 'edited_page.html';
  a.click();
}}

function finishEditing(){{
  hideError();
  downloadWord().then(function(){{
    setTimeout(downloadHtml, 400);
    setTimeout(function(){{ openOverlay('finish_overlay'); }}, 500);
  }});
}}
</script>
</body>
</html>"""


def fetch_html(url: str, dynamic: bool) -> str:
    if dynamic:
        from playwright.sync_api import sync_playwright
        with sync_playwright() as p:
            browser = p.chromium.launch()
            page = browser.new_page()
            page.goto(url, wait_until="networkidle", timeout=30000)
            html = page.content()
            browser.close()
        return html
    import requests
    resp = requests.get(url, timeout=20, headers={"User-Agent": "EditorTool/1.0"})
    resp.raise_for_status()
    return resp.text


def absolutize(soup: BeautifulSoup, base_url: str) -> None:
    for tag, attr in [("img", "src"), ("a", "href"), ("link", "href"), ("source", "src")]:
        for el in soup.find_all(tag):
            if el.get(attr):
                el[attr] = urljoin(base_url, el[attr])


def has_ignored_hint(el) -> bool:
    identifier = " ".join([el.get("id", "")] + el.get("class", [])).lower()
    return any(h in identifier for h in IGNORED_CLASS_ID_HINTS)


def is_inside_heading(node) -> bool:
    parent = node.parent
    while parent is not None:
        if getattr(parent, "name", None) in ("h1", "h2", "h3", "h4", "h5", "h6"):
            return True
        parent = parent.parent
    return False


def is_editable_text(node) -> bool:
    if not isinstance(node, NavigableString):
        return False
    if not str(node).strip():
        return False
    if node.parent and getattr(node.parent, "name", None) == "img":
        return False

    inside_heading = is_inside_heading(node)

    parent = node.parent
    while parent is not None:
        tag = getattr(parent, "name", None)
        if tag in HARD_REMOVE_TAGS or tag in ("img", "svg"):
            return False
        if not inside_heading and tag in SOFT_IGNORE_TAGS:
            return False
        if not inside_heading and has_ignored_hint(parent):
            return False
        parent = parent.parent
    return True


def make_editable_html(url: str, dynamic: bool) -> str:
    html = fetch_html(url, dynamic)
    soup = BeautifulSoup(html, "html.parser")
    for bad in soup(list(HARD_REMOVE_TAGS)):
        bad.decompose()
    absolutize(soup, url)

    counter = 0
    for text_node in list(soup.find_all(string=True)):
        if is_editable_text(text_node):
            span = soup.new_tag("span")
            span["data-edit-id"] = f"t{counter}"
            span["contenteditable"] = "true"
            span["style"] = "outline:1px dashed rgba(230,0,126,.4);outline-offset:1px;"
            span.string = str(text_node)
            text_node.replace_with(span)
            counter += 1

    init_script = soup.new_tag("script")
    init_script.string = """
    window.__ORIGINALS__ = {};
    document.querySelectorAll('[data-edit-id]').forEach(function(el){
      var id = el.getAttribute('data-edit-id');
      window.__ORIGINALS__[id] = el.textContent;
    });
    """
    if soup.body:
        soup.body.append(init_script)
    else:
        soup.append(init_script)

    return str(soup)


@app.route("/")
def index():
    url = request.args.get("url", "")
    dynamic = request.args.get("dynamic", "false") == "true"

    if not url:
        return STEP1_SHELL

    try:
        iframe_srcdoc = make_editable_html(url, dynamic).replace('"', "&quot;")
    except Exception as e:
        iframe_srcdoc = f"<p style='padding:20px;color:red'>Error: {e}</p>".replace('"', "&quot;")

    return STEP2_SHELL.format(url_value=url, iframe_srcdoc=iframe_srcdoc)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text or "")
    run.font.size = Pt(10)
    run.bold = bold
    if color:
        run.font.color.rgb = color


@app.route("/export/docx", methods=["POST"])
def export_docx():
    payload = request.get_json()
    blocks = payload.get("blocks", [])

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height

    doc.add_heading("Content Update", level=1)
    doc.add_paragraph(
        "Left column: current content on the site. Right column: new content to paste into the CMS. "
        "Only rows with actual changes are shown."
    )

    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "Current content", bold=True)
    set_cell_text(hdr[1], "New content", bold=True)

    changed_color = RGBColor(0x00, 0x72, 0xCE)
    changes_found = 0

    for b in blocks:
        original = (b.get("original") or "").strip()
        current = (b.get("current") or "").strip()
        if not original and not current:
            continue
        if original == current:
            continue
        changes_found += 1
        row = table.add_row().cells
        set_cell_text(row[0], original)
        set_cell_text(row[1], current, bold=True, color=changed_color)

    if changes_found == 0:
        row = table.add_row().cells
        set_cell_text(row[0], "No changes detected.")
        set_cell_text(row[1], "")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name="content_update.docx",
    )


if __name__ == "__main__":
    app.run(debug=True, port=5000)
